import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO
import openai

# Set your OpenAI API key
openai.api_key = st.secrets.get("openai_api_key", "your-openai-key-here")

st.set_page_config(page_title="📝 Response to Reviewers", layout="wide")
st.title("📝 Response to Reviewers Tool")
st.markdown("This tool helps researchers generate structured and well-formatted responses to peer reviewer comments.")

# Journal and Manuscript Info
journal_name = st.text_input("📘 Journal's Name")
manuscript_id = st.text_input("🆔 Manuscript ID")
manuscript_title = st.text_input("📝 Manuscript Title")

st.markdown("#### 📄 Cover Note (Automatically added)")
st.markdown("""We would like to thank you for your valuable time and insightful comments to improve the quality of this manuscript. We have now addressed these comments and revised the manuscript thoroughly. Please see our detailed responses and revisions below.

Thanks  
Authors""")

st.divider()

# Initialize reviewer state
if "reviewers" not in st.session_state:
    st.session_state.reviewers = [{"id": 1, "comments": [{"comment": "", "response": ""}]}]

# Functions
def add_reviewer():
    st.session_state.reviewers.append({
        "id": len(st.session_state.reviewers) + 1,
        "comments": [{"comment": "", "response": ""}]
    })

def add_comment(reviewer_index):
    st.session_state.reviewers[reviewer_index]["comments"].append({"comment": "", "response": ""})

def generate_ai_response(comment_text):
    prompt = f"""You're helping a researcher respond to a peer reviewer. Write a formal, appreciative, and specific reply to this reviewer comment:

Comment: {comment_text}
Response:"""
    try:
        res = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=200
        )
        return res.choices[0].message.content.strip()
    except Exception as e:
        return f"[Error generating AI response: {e}]"

# Reviewer Sections
for i, reviewer in enumerate(st.session_state.reviewers):
    st.subheader(f"👤 Reviewer {reviewer['id']}")
    for j, pair in enumerate(reviewer["comments"]):
        comment = st.text_area(f"Comment {j+1}", key=f"comment_{i}_{j}")
        ai_button = st.button("💡 Generate AI Response", key=f"ai_button_{i}_{j}")
        if ai_button and comment:
            ai_response = generate_ai_response(comment)
            st.session_state.reviewers[i]["comments"][j]["response"] = ai_response
        response = st.text_area(
            f"Response {j+1}",
            value=st.session_state.reviewers[i]["comments"][j]["response"],
            key=f"response_{i}_{j}"
        )
        reviewer["comments"][j]["comment"] = comment
        reviewer["comments"][j]["response"] = response
    st.button("➕ Add Comment", on_click=add_comment, args=(i,), key=f"add_comment_{i}")
    st.markdown("---")

st.button("➕ Add Reviewer", on_click=add_reviewer)

# Generate Word File
def generate_docx():
    doc = Document()
    doc.add_heading(journal_name or "Journal Name", level=1)
    doc.add_paragraph(f"Manuscript ID: {manuscript_id}")
    doc.add_paragraph(f"Manuscript Title: {manuscript_title}")
    doc.add_paragraph(
        "We would like to thank you for your valuable time and insightful comments to improve the quality of this manuscript. "
        "We have now addressed these comments and revised the manuscript thoroughly. Please see our detailed responses and revisions below.\n\nThanks,\nAuthors"
    )

    for reviewer in st.session_state.reviewers:
        doc.add_heading(f"Reviewer {reviewer['id']}", level=2)
        for idx, pair in enumerate(reviewer["comments"], 1):
            doc.add_paragraph(f"Comment {idx}:", style='List Number')
            p1 = doc.add_paragraph(pair["comment"])
            p1.paragraph_format.left_indent = Pt(20)
            doc.add_paragraph("Response:", style='List Number')
            p2 = doc.add_paragraph(pair["response"])
            p2.paragraph_format.left_indent = Pt(20)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("📄 Generate Word File"):
    buffer = generate_docx()
    st.download_button(
        label="📥 Download Response_to_Reviewers.docx",
        data=buffer,
        file_name="Response_to_Reviewers.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
